import requests
import xml.etree.ElementTree as ET
import os
from string import punctuation
import re
import csv
import logging
import hashlib
from functools import reduce
from docx import Document

logging.basicConfig(filename='logs/paper-kinderformularium.log',
                    format='%(asctime)s|%(levelname)-8s|%(message)s',
                    level=logging.INFO,
                    datefmt='%Y-%m-%d %H:%M:%S')

def hash_shingle(shingle):
    """Hash a shingle to an integer using SHA-1."""
    return int(hashlib.sha1(shingle.encode('utf-8')).hexdigest(), 16)

def jaccard_similarity(s1, s2):
    """Compute the Jaccard similarity between two input texts after shingling & hashing."""
    shingle1 = {s1[i:i+3] for i in range(len(s1) - 3 + 1)}  # shingle size: 3
    shingle2 = {s2[i:i+3] for i in range(len(s2) - 3 + 1)}  # shingle size: 3

    minhashes1, minhashes2 = [], []
    for i in range(200):
        # Use different hash functions by XORing with an incremental value
        hash_fn = lambda x: hash_shingle(x) ^ i
        min_hash1 = min(map(hash_fn, shingle1))
        min_hash2 = min(map(hash_fn, shingle2))
        minhashes1.append(min_hash1)
        minhashes2.append(min_hash2)

    return sum(1 for i in range(len(minhashes1)) if minhashes1[i] == minhashes2[i]) / len(minhashes1)

def search_API(query):
    """Retrieves ids of relevant papers from PubMed API

    Parameters
    ----------
    query : string
        DOI or title from a targeted paper

    Returns
    -------
    results : list of integers
        Plain list of paper ids
    """
    url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
    params = {
        "db": 'pubmed',
        "term": query,
        "sort": 'relevance',
        "retmax": '20',
        "retmode": 'xml'
    }
    response = requests.get(url, params=params)
    
    root = ET.fromstring(response.text)
    ids = [id_elem.text for id_elem in root.findall(".//Id")]
    return ids

def fetch_details(id_list):
    """Retrieves the papers from PubMed API by their ids

    Parameters
    ----------
    id_list : list of integers

    Returns
    -------
    results : list of dictionaries
    """
    url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi"
    params = {
        "db": 'pubmed',
        "id": ",".join(id_list),
        "retmode": 'xml'
    }
    response = requests.get(url, params=params)

    root = ET.fromstring(response.text)
    articles = []
    for article in root.findall(".//PubmedArticle"):
        title_elem = article.find(".//ArticleTitle")
        abstract_elem = article.find(".//Abstract/AbstractText")
        doi_elem = article.find(".//ArticleIdList/ArticleId[@IdType='doi']")

        title = title_elem.text if title_elem is not None else "No title"
        abstract = abstract_elem.text if abstract_elem is not None else "No abstract"
        doi = doi_elem.text if doi_elem is not None else "No DOI"

        articles.append({"title": title, "abstract": abstract, "doi": doi})
    return articles

def create_csv_path(rba_path):
    """Creates a csv path for the output

    Parameters
    ----------
    rba_path : string
        Path to the file to be processed.
        This file must be a .docx formatted file.

    Returns
    -------
    csv_path : string
    """
    csv_dir = os.path.abspath(os.path.join(os.path.dirname(rba_path), 'csv'))
    if not os.path.exists(csv_dir):
        os.makedirs(csv_dir)

    head, filename = os.path.split(rba_path)
    no = re.findall("^([0-9]+)([^.]+)(.*\.docx)", filename)
    if no and no[0][1] != '':
        csv_path = os.path.join(csv_dir, (no[0][0] + no[0][2]).replace('.docx', '.csv'))
    else:
        csv_path = os.path.join(csv_dir, filename.replace('.docx', '.csv'))

    return csv_path

def collectFromEndnote(rba_path):
    """Extracts the reference list from a file with Endnote

    Parameters
    ----------
    rba_path : string
        Path to the file to be processed.
        This file must be a .docx formatted file.
        The file must contain a text sayin at least "referen"
        in order to function properly

    Returns
    -------
    references_list : list of dictionaries
        Contains the extracted DOI, authors, and title
        for each reference as an object.
    """
    doc = Document(rba_path) # create an instance of a word document we want to open

    findBegin = False
    references_list = []

    for i in range(0, len(doc.paragraphs)):
        if (doc.paragraphs[i].text.lower().startswith('references') or \
            doc.paragraphs[i].text.lower().startswith('referen')):
            for run in doc.paragraphs[i].runs:
                if run.bold:
                    findBegin = True

        if (findBegin == True):
            ref = doc.paragraphs[i].text.strip()
            if (ref.lower().startswith('referen')):
                continue
            elif (ref.strip() == ''):
                continue
            else:

                search_item = {
                    'original_text': ref,
                    'p_authors': '',
                    'p_title': '',
                    'p_doi': ''
                }

                x1 = re.findall("([^.]+).([^?|^.]+).", ref)
                x2 = re.findall("(.*)\d{4}.?\s?(\".*\")", ref)
                x3 = re.findall(r'doi:?\s?(.+).?', ref)
                x4 = re.findall("(.*et al.?)([^?|^.]+).", ref)

                if x3:
                    search_item['p_doi'] = x3[0].strip(punctuation).strip()

                if x1:
                    search_item['p_authors'] = x1[0][0].strip(punctuation).strip()
                    search_item['p_title'] = x1[0][1].strip(punctuation).strip()

                if x2:
                    search_item['p_authors'] = x2[0][0].strip(punctuation).strip()
                    search_item['p_title'] = x2[0][1].strip(punctuation).strip()

                if ('et al.' in ref) and x4:
                    search_item['p_authors'] = x4[0][0].strip(punctuation).strip()
                    search_item['p_title'] = x4[0][1].strip(punctuation).strip()

                elif (ref.lower().startswith('who.') or ref.lower().startswith('lci.') \
                      or ref.lower().startswith('nvn.') or ref.lower().startswith('nice.')):
                    continue


            references_list.append(search_item)
    return references_list

def collectFromTables(rba_path):
    """Extracts the reference list from a file with a table

    Parameters
    ----------
    rba_path : string
        Path of the file to be processed

    Returns
    -------
    references_list : list of dictionaries
        Contains the extracted authors and title for each reference as an object
    """
    doc = Document(rba_path)

    data = []
    references_list = []
    for table in doc.tables:

        if len(table.rows) == 0:
            continue

        for i, row in enumerate(table.rows):

            if len(row.cells) == 0:
                continue

            for j, cell in enumerate(row.cells):

                if (len(cell.paragraphs) > 0
                    and len(cell.paragraphs[0].runs) > 0
                    and cell.paragraphs[0].runs[0].bold
                    and ('samenvatting' in cell.text.lower()
                         or 'summary' in cell.text.lower())):

                    bold_text = ''
                    p_title = ''
                    p_authors = ''

                    for p, paragraph in enumerate(cell.paragraphs):

                        if len(paragraph.runs) < 1:
                            continue

                        if (p == 0
                            and paragraph.runs[0].bold
                            and not paragraph.runs[-1].bold):

                            for r, run in enumerate(paragraph.runs):

                                if run.bold:
                                    bold_text += run.text.strip() + ' '

                                else:
                                    p_title = bold_text

                                if p_title and not run.bold and run.text != ' ':

                                    p_authors += run.text.strip() + ' '

                        elif (paragraph.runs[0].bold
                              and paragraph.runs[-1].bold and p_title == ''):

                            bold_text += paragraph.text.strip() + ' '

                        elif (not paragraph.runs[-1].bold
                              and p > 0
                              and len(cell.paragraphs[p-1].runs) > 0
                              and cell.paragraphs[p-1].runs[-1].bold):

                            p_title = bold_text

                            p_authors += paragraph.text.strip() + ' '

                        elif (paragraph.runs[0].underline
                              and (paragraph.runs[0].text.lower().startswith('samenvatting')
                                   or paragraph.runs[0].text.lower().startswith('summary'))):
                            break

                    if p_title.strip() != ''  and re.search(r'^[0-9].[0-9]', p_title) is None:
                        references_list.append({
                            'p_title': p_title.replace('\xa0',' ').replace('\n', '').strip(),
                            'p_authors': p_authors.replace('\xa0',' ').replace('\n', '').strip(),
                        })

    return references_list

def pubmed2csv(references_list, csv_path):
    """Creates a csv file consists of the references

    Parameters
    ----------
    references_list : list of dictionaries
        Contains dictionaries with paper DOIs, authors and titles

    csv_path : string
        corresponding csv path for the file

    Returns
    -------
    Creates a csv file on the given CSV path
    """
    logging.info('PROCESSING THE FILE: {}'.format(csv_path))

    # print to console
    print(f"Processing {len(references_list)} references...")

    if not os.path.exists(csv_path):
        pmids = []
        header = ['pubmed_id', 'title', 'abstract', 'doi', 'final_included']
        with open(csv_path, 'w', newline='', encoding='ISO-8859-15', errors='ignore') as csvfile:
            cw = csv.writer(csvfile, delimiter=',')
            cw.writerow(header)

    elif os.path.exists(csv_path):
        with open(csv_path, 'r', newline='', encoding='ISO-8859-15', errors='ignore') as csvfile:
            cw = csv.reader(csvfile)
            header = next(cw)

            rows = []
            for row in cw:
                rows.append(row)

        pmids = [l[0] for l in rows]

    #Loop over the reference list
    ref_num = 0
    for ref in references_list:
        ref_num += 1
        print(f"Processing reference {ref_num}/{len(references_list)}...")

        has_doi = 'p_doi' in ref and ref['p_doi'] != ''
        search_query = ref['p_doi'] if has_doi else ref['p_title']

        # No p_doi and no p_title
        if search_query is None:
            continue

        try:
            id_list = []
            # Try with doi first, if it is there
            if has_doi:
                id_list = search_API(search_query)

            # Try again with title if doi did not work or doi not provided
            if id_list == []:
                search_query = ref['p_title']
                id_list = search_API(search_query)

            # Try one more time, this time removing the last section in case journal title in title obfuscates
            # In case p_title == "{title}.{journal}."
            if id_list == [] and search_query.count('.') > 1:
                search_query = search_query.split('.')[0]
                id_list = search_API(search_query)

            # If no results found in PubMed, skip and log
            if not any(id_list):
                logging.warning(f"!! The PubMed search for {search_query} did not return any results! \
                                Paper title is {ref['p_title']}!")
                continue

            papers = fetch_details(id_list)
            referenceFound = False
            for paperIndex, paper in enumerate(papers):
                paperTitle = paper['title']
                jaccard_score = jaccard_similarity(paperTitle.strip(punctuation), ref['p_title'].strip(punctuation))
                if (jaccard_score >= 0.9):
                    if ((id_list[paperIndex] not in pmids) or pmids == []):
                        data = [id_list[paperIndex], paperTitle, paper['abstract'], paper['doi'], 1]
                        with open(csv_path, 'a', newline='', encoding='ISO-8859-15', errors='ignore') as csvfile:
                            cw = csv.writer(csvfile, delimiter=',')
                            cw.writerow(data)
                        referenceFound = True
                        pmids.append(id_list[paperIndex])
                        break
                    elif (id_list[paperIndex] in pmids):
                        logging.debug('The reference "{}" is already in the csv file, therefore skipped.'.format(ref['p_title']))
                        referenceFound = True
                        break
            if (not referenceFound) and (id_list[paperIndex] not in pmids):
                logging.error('!!!!! The reference could not be retrieved from the PubMed database. Search query was: "{}". Paper title was: "{}"'.format(search_query, ref['p_title']))
        except Exception as err:
            logging.error('!!! The reference could not be found automaticaly for the title "{}". Error message: {}'.format(ref['p_title'], err))

if __name__ == "__main__":
    rba_path = os.path.abspath('docs/3b. Risicoanalyse kinderformularium clonazepam epilepsie.docx')
    csv_path = create_csv_path(rba_path)

    # Collect the Endnote references from the RBA file
    references_list_endnote = collectFromEndnote(rba_path)

    # Write the extracted references to the output file
    pubmed2csv(references_list_endnote, csv_path)

    # Collect the table references from the RBA file
    references_list_table = collectFromTables(rba_path)

    # Write the extracted references to the output file
    pubmed2csv(references_list_table, csv_path)
